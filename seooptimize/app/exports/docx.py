"""DOCX Export Engine."""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.logging import get_logger
from app.models.page import PageData, ScoreStatus
from app.models.project import ProjectConfig

logger = get_logger(__name__)


def generate_docx_report(project: ProjectConfig, pages: list[PageData]) -> bytes:
    """Generate a DOCX SEO report.

    Args:
        project: Project configuration.
        pages: List of analysed pages.

    Returns:
        DOCX file as bytes.
    """
    doc = Document()

    # ── Document styles ───────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Title page ────────────────────────────────────────────────────────────
    title = doc.add_heading("SEOOptimize — SEO Audit Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Business: {project.business_name}")
    doc.add_paragraph(f"City: {project.target_city}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph(f"Website: {project.website_url}")
    doc.add_paragraph("")

    # ── Site summary ──────────────────────────────────────────────────────────
    scores = [int(p.scores.total) for p in pages if p.extraction_complete]
    avg_score = sum(scores) // len(scores) if scores else 0

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(f"Overall Site Score: {avg_score}/100")
    doc.add_paragraph(f"Pages Analysed: {len(pages)}")

    # Six-axis table (first page)
    if pages:
        page = pages[0]
        doc.add_heading("Scoring Axes", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Axis"
        hdr[1].text = "Weight"
        hdr[2].text = "Score"

        axes = [
            ("Local SEO", "30%", page.scores.local_seo, 30),
            ("Content Quality", "25%", page.scores.content_quality, 25),
            ("Technical SEO", "15%", page.scores.technical_seo, 15),
            ("Conversion Signals", "15%", page.scores.conversion_signals, 15),
            ("On-Page Metadata", "10%", page.scores.on_page_metadata, 10),
            ("Competitor Gap", "5%", page.scores.competitor_gap, 5),
        ]
        for name, weight, val, max_v in axes:
            row = table.add_row().cells
            row[0].text = name
            row[1].text = weight
            row[2].text = f"{val:.0f}/{max_v}"

    # ── Per-page sections ─────────────────────────────────────────────────────
    for page in pages:
        doc.add_page_break()
        short_url = page.url.replace("https://", "").replace("http://", "")
        doc.add_heading(f"Page: {short_url}", level=1)
        doc.add_paragraph(f"Score: {int(page.scores.total)}/100")

        # Field audit table
        doc.add_heading("SEO Field Audit", level=2)
        field_table = doc.add_table(rows=1, cols=3)
        field_table.style = "Table Grid"
        hdr = field_table.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Status"
        hdr[2].text = "Note"

        fields = [
            ("Meta Title", page.extracted.meta_title),
            ("Meta Description", page.extracted.meta_description),
            ("H1 Tag", page.extracted.h1),
            ("H2 Structure", page.extracted.h2_structure),
            ("Word Count", page.extracted.word_count),
            ("Images Alt Text", page.extracted.images_with_alt),
            ("Phone Above Fold", page.extracted.phone_above_fold),
            ("Schema Markup", page.extracted.schema_markup),
            ("Canonical Tag", page.extracted.canonical_tag),
            ("Mobile Viewport", page.extracted.mobile_viewport),
            ("NAP on Page", page.extracted.nap_on_page),
            ("Internal Links", page.extracted.internal_links),
            ("Page Load Time", page.extracted.page_load_time),
            ("HTTPS", page.extracted.https),
        ]

        for label, field in fields:
            row = field_table.add_row().cells
            row[0].text = label
            row[1].text = field.status.value.upper()
            row[2].text = (field.note or "")[:100]

        # Recommendations
        cards_data = page.ai_analysis.get("recommendation_cards", [])
        if cards_data:
            doc.add_heading("AI Recommendations", level=2)
            for card in cards_data[:10]:
                priority = card.get("priority", "warning").upper()
                label = card.get("label", "")
                problem = card.get("problem", "")
                fix = card.get("suggested_fix", "")

                para = doc.add_paragraph()
                run = para.add_run(f"[{priority}] {label}")
                run.bold = True
                doc.add_paragraph(f"Problem: {problem}")
                doc.add_paragraph(f"Suggested Fix: {fix[:200]}")
                doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
